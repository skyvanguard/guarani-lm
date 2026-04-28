"""
Scraper for Oremba'e (orembae.org.py) - Biblioteca Virtual de Literatura Guaraní.

Sources:
  1. API at articulos.instituto.org.py/api/libros — 13,469+ poems with PDF files
  2. Static DOC/DOCX poems at orembae.org.py/assets/poems/

Filters by idioma=="Guaraní" to get only Guaraní-language content.
Downloads PDFs and extracts text using pdfplumber.
Downloads DOC/DOCX files and extracts text using python-docx.
"""
import json
import os
import re
import sys
import time
import argparse
import urllib.parse
from pathlib import Path

import requests

# Base directories
SCRIPT_DIR = Path(__file__).parent
PROJECT_DIR = SCRIPT_DIR.parent
OUTPUT_DIR = PROJECT_DIR / "data" / "raw" / "orembae"
PDF_DIR = OUTPUT_DIR / "pdfs"
DOC_DIR = OUTPUT_DIR / "docs"
TEXT_DIR = OUTPUT_DIR / "texts"

# API endpoints
API_BASE = "https://articulos.instituto.org.py/api"
LIBROS_URL = f"{API_BASE}/libros"
ARCHIVO_URL = f"{API_BASE}/archivo"

# Static poems base
POEMS_BASE = "https://www.orembae.org.py/assets/poems"

# Known static poem files (Guaraní = (G) in filename)
STATIC_POEMS = [
    "1.16. Nde Che morena mi (G).docx",
    "1.28. A Emiliano R (G).docx",
    "10.12. Pee Soldados añete pe (G).doc",
    "13.9. Jardín Laurel ty (G).rtf",
    "14.14. Morenita (G).doc",
    "14.40. Cuña Jha yvoty (G).docx",
    "15.7. Mutilado Recové (G).docx",
    "2.13. Cuarahy reitke (G).doc",
    "2.8. Itaugueña (G).doc",
    "20.26. Cuña Pynandi (G).docx",
    "20.45. Pirayú poty (G).docx",
    "21.19. Musiqueada Jazmin guipe (G).docx",
    "23.2. Voz de ayer (G).docx",
    "23.20. Asunción del Paraguay FALTA AUTOR Y FECHA (G).docx",
    "24.33. Ore rera cuera (G).docx",
    "25.2. Tyvytá Yasy pîajhù (G).docx",
    "25.6. Pyjharé Marangatú (G).docx",
    "27.13. Nde rera mi (G).docx",
    "27.31. Mbocaya poty (G).docx",
    "27.44. TOPYTU_U CHE RECOVE (G).docx",
    "31.26. Peeme che reindy cuera (G).docx",
    "31.8. Pycazu (G).docx",
    "5.13. Tupãsy Caacupé (G).doc.docx",
    "5.9. Ñande retá ñe_e (G).doc",
]

SESSION = requests.Session()
SESSION.headers.update({
    "User-Agent": "GuaraniLM-DataCollector/1.0 (Academic Research; guarani-lm project)"
})


def fetch_all_guarani_metadata(limit=100, max_pages=None):
    """Fetch all book metadata from the API, filtering for Guaraní."""
    all_books = []
    page = 1

    while True:
        if max_pages and page > max_pages:
            break

        print(f"  Fetching page {page} (limit={limit})...")
        try:
            resp = SESSION.get(LIBROS_URL, params={"page": page, "limit": limit}, timeout=30)
            resp.raise_for_status()
            data = resp.json()
        except Exception as e:
            print(f"  ERROR fetching page {page}: {e}")
            break

        books = data.get("data", [])
        meta = data.get("meta", {})

        if not books:
            break

        # Filter Guaraní books
        guarani_books = [b for b in books if b.get("idioma", "").lower() == "guaraní"]
        all_books.extend(guarani_books)

        total = meta.get("total_items", 0)
        total_pages = meta.get("total_pages", 0)
        has_more = meta.get("has_more", False)

        print(f"    Found {len(guarani_books)}/{len(books)} Guaraní books "
              f"(total so far: {len(all_books)}, page {page}/{total_pages})")

        if not has_more:
            break

        page += 1
        time.sleep(0.5)  # Be polite

    return all_books


def download_pdf(archivo, dest_dir):
    """Download a PDF from the API."""
    url = f"{ARCHIVO_URL}/{archivo}"
    dest = dest_dir / archivo

    if dest.exists():
        return dest

    try:
        resp = SESSION.get(url, timeout=60)
        resp.raise_for_status()
        dest.write_bytes(resp.content)
        return dest
    except Exception as e:
        print(f"    ERROR downloading {archivo}: {e}")
        return None


def download_static_poem(filename, dest_dir):
    """Download a static poem file (DOC/DOCX/RTF)."""
    encoded = urllib.parse.quote(filename)
    url = f"{POEMS_BASE}/{encoded}"
    dest = dest_dir / filename

    if dest.exists():
        return dest

    try:
        resp = SESSION.get(url, timeout=30)
        resp.raise_for_status()
        dest.write_bytes(resp.content)
        return dest
    except Exception as e:
        print(f"    ERROR downloading {filename}: {e}")
        return None


def extract_text_from_pdf(pdf_path):
    """Extract text from PDF using pdfplumber."""
    try:
        import pdfplumber
        text_parts = []
        with pdfplumber.open(pdf_path) as pdf:
            for page in pdf.pages:
                page_text = page.extract_text()
                if page_text:
                    text_parts.append(page_text.strip())
        return "\n\n".join(text_parts)
    except Exception as e:
        print(f"    ERROR extracting text from {pdf_path.name}: {e}")
        return ""


def extract_text_from_docx(docx_path):
    """Extract text from DOCX using python-docx."""
    try:
        import docx
        doc = docx.Document(str(docx_path))
        paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
        return "\n".join(paragraphs)
    except Exception as e:
        print(f"    ERROR extracting text from {docx_path.name}: {e}")
        return ""


def extract_text_from_doc(doc_path):
    """Extract text from DOC using antiword (if available) or textract."""
    # Try with python-docx first (works for some .doc files)
    try:
        import docx
        doc = docx.Document(str(doc_path))
        paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
        if paragraphs:
            return "\n".join(paragraphs)
    except Exception:
        pass

    # Fallback: read raw bytes and extract text heuristically
    try:
        raw = doc_path.read_bytes()
        # Try to find text between common markers in OLE2 format
        text = raw.decode("latin-1", errors="ignore")
        # Remove control characters but keep newlines
        text = re.sub(r'[\x00-\x08\x0b\x0c\x0e-\x1f]', '', text)
        # Extract readable portions
        readable = re.findall(r'[\w\s\.,;:!?¡¿áéíóúñãĩỹÁÉÍÓÚÑÃĨỸ]{20,}', text)
        if readable:
            return "\n".join(readable)
    except Exception:
        pass

    return ""


def main():
    parser = argparse.ArgumentParser(description="Scrape Oremba'e Guaraní literature")
    parser.add_argument("--max-pages", type=int, default=None,
                        help="Max API pages to fetch (default: all)")
    parser.add_argument("--skip-pdfs", action="store_true",
                        help="Skip PDF downloads")
    parser.add_argument("--skip-static", action="store_true",
                        help="Skip static poem downloads")
    parser.add_argument("--limit", type=int, default=100,
                        help="Items per API page (default: 100)")
    args = parser.parse_args()

    # Create directories
    for d in [OUTPUT_DIR, PDF_DIR, DOC_DIR, TEXT_DIR]:
        d.mkdir(parents=True, exist_ok=True)

    # === Phase 1: API metadata ===
    print("=" * 60)
    print("PHASE 1: Fetching Guaraní book metadata from API")
    print("=" * 60)

    metadata_file = OUTPUT_DIR / "guarani_metadata.jsonl"

    books = fetch_all_guarani_metadata(limit=args.limit, max_pages=args.max_pages)
    print(f"\nTotal Guaraní books found: {len(books)}")

    with open(metadata_file, "w", encoding="utf-8") as f:
        for book in books:
            f.write(json.dumps(book, ensure_ascii=False) + "\n")
    print(f"Metadata saved to {metadata_file}")

    # === Phase 2: Download PDFs ===
    if not args.skip_pdfs:
        print("\n" + "=" * 60)
        print(f"PHASE 2: Downloading {len(books)} Guaraní PDFs")
        print("=" * 60)

        downloaded = 0
        for i, book in enumerate(books):
            archivo = book.get("archivo", "")
            if not archivo:
                continue

            print(f"  [{i+1}/{len(books)}] {book.get('titulo', 'Unknown')} -> {archivo}")
            result = download_pdf(archivo, PDF_DIR)
            if result:
                downloaded += 1

            if (i + 1) % 50 == 0:
                time.sleep(1)  # Rate limiting

        print(f"\nDownloaded {downloaded}/{len(books)} PDFs")

    # === Phase 3: Download static poems ===
    if not args.skip_static:
        print("\n" + "=" * 60)
        print(f"PHASE 3: Downloading {len(STATIC_POEMS)} static Guaraní poems")
        print("=" * 60)

        for i, filename in enumerate(STATIC_POEMS):
            print(f"  [{i+1}/{len(STATIC_POEMS)}] {filename}")
            download_static_poem(filename, DOC_DIR)
            time.sleep(0.3)

    # === Phase 4: Extract text ===
    print("\n" + "=" * 60)
    print("PHASE 4: Extracting text from downloaded files")
    print("=" * 60)

    all_texts = []

    # Extract from PDFs
    pdf_files = list(PDF_DIR.glob("*.pdf"))
    print(f"\nExtracting text from {len(pdf_files)} PDFs...")
    for i, pdf_path in enumerate(pdf_files):
        text = extract_text_from_pdf(pdf_path)
        if text and len(text.strip()) > 20:
            all_texts.append({
                "source": "orembae_api",
                "file": pdf_path.name,
                "text": text.strip(),
                "chars": len(text.strip()),
            })
        if (i + 1) % 100 == 0:
            print(f"    Processed {i+1}/{len(pdf_files)} PDFs...")

    # Extract from DOC/DOCX
    doc_files = list(DOC_DIR.glob("*"))
    print(f"\nExtracting text from {len(doc_files)} DOC/DOCX files...")
    for doc_path in doc_files:
        if doc_path.suffix.lower() == ".docx":
            text = extract_text_from_docx(doc_path)
        elif doc_path.suffix.lower() == ".doc":
            text = extract_text_from_doc(doc_path)
        else:
            continue

        if text and len(text.strip()) > 20:
            all_texts.append({
                "source": "orembae_static",
                "file": doc_path.name,
                "text": text.strip(),
                "chars": len(text.strip()),
            })

    # Save extracted texts
    output_file = OUTPUT_DIR / "guarani_texts.jsonl"
    with open(output_file, "w", encoding="utf-8") as f:
        for item in all_texts:
            f.write(json.dumps(item, ensure_ascii=False) + "\n")

    total_chars = sum(t["chars"] for t in all_texts)
    print(f"\n{'=' * 60}")
    print(f"RESULTS:")
    print(f"  Total texts extracted: {len(all_texts)}")
    print(f"  Total characters: {total_chars:,}")
    print(f"  Estimated tokens: ~{total_chars // 4:,}")
    print(f"  Output: {output_file}")
    print(f"{'=' * 60}")


if __name__ == "__main__":
    main()
